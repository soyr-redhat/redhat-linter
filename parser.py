import os
from docx import Document
from typing import List, Dict
from docling.document_converter import DocumentConverter

class RedHatParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        
    def get_structured_content(self) -> List[Dict[str, str]]:
        """
        Parses the docx and returns a list of dictionaries containing 
        the text and its 'type' (heading, body, list_item).
        """
        structured_data = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Determine the 'type' based on docx style
            # This helps the AI apply formatting rules vs tone rules
            style = para.style.name.lower()
            
            content_type = "body"
            if "heading" in style:
                content_type = "heading"
            elif "list" in style or text.startswith(('•', '-', '*')):
                content_type = "list_item"
                
            structured_data.append({
                "text": text,
                "type": content_type,
                "style": style
            })
            
        return structured_data

    def extract_full_text(self) -> str:
        """Utility for a quick overall dump if needed."""
        return "\n".join([p.text for p in self.doc.paragraphs])

# --- Logic for the 'Guides' Directory ---

def process_document_with_docling(file_path: str) -> str:
    """
    Uses docling to process various document formats (PDF, DOCX, HTML, etc.)
    and extract clean text content.
    """
    try:
        converter = DocumentConverter()
        result = converter.convert(file_path)
        return result.document.export_to_markdown()
    except Exception as e:
        return f"Error processing {file_path} with docling: {str(e)}"

def load_guides(guides_dir: str = "guides") -> Dict[str, str]:
    """
    Reads all document files in the guides directory using docling.
    Supports: .md, .pdf, .docx, .html, .htm, .txt
    """
    guides_context = {}
    if not os.path.exists(guides_dir):
        os.makedirs(guides_dir)
        return {"error": "Guides directory was missing and has been created."}

    supported_extensions = ('.md', '.pdf', '.docx', '.html', '.htm', '.txt')

    for filename in os.listdir(guides_dir):
        file_path = os.path.join(guides_dir, filename)

        if filename.lower().endswith(supported_extensions):
            guide_name = os.path.splitext(filename)[0]

            # Use docling for all formats except plain markdown
            if filename.endswith('.md'):
                with open(file_path, "r", encoding="utf-8") as f:
                    guides_context[guide_name] = f.read()
            else:
                guides_context[guide_name] = process_document_with_docling(file_path)

    return guides_context

# Example usage for testing:
if __name__ == "__main__":
    # parser = RedHatParser("your_draft.docx")
    # content = parser.get_structured_content()
    # print(f"Parsed {len(content)} elements.")
    
    guides = load_guides()
    print(f"Loaded {len(guides)} style guides.")
